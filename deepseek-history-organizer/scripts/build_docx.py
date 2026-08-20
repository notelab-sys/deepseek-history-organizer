#!/usr/bin/env python3
"""Convert Markdown file(s) to a Word (.docx) document.

Typography follows a Chinese journal style layout:
- Body: Times New Roman (Latin) + SimSun (Chinese), 10.5 pt, line spacing 1.3,
  paragraphs start with a two-character first-line indent.
- Title: SimHei 20 pt centered; H2 sections 14 pt SimSun; H3+ SimHei.
- Gene symbols and Latin binomials are italic; protein/enzyme abbreviations
  stay regular.
- Reference-list entries use a hanging indent at 8 pt with 1.15 line spacing.
- Centered footer page numbers: 第 X 页 / 共 Y 页.

With --lang en, the layout switches to an English journal style:
- Body: Times New Roman 12 pt, line spacing 1.5, 0.5-inch first-line indent.
- Title: Times New Roman bold 16 pt centered; H2 bold 14 pt; H3+ bold 12 pt.
- Reference-list entries use a 0.5-inch hanging indent at 10 pt.
- Footer page numbers: Page X of Y.

Multiple Markdown inputs are merged in order (e.g. content + references appendix).

Usage:
  python build_docx.py <input.md> [...] [-o output.docx] [--lang zh|en]
"""

import argparse
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

EAST = "SimSun"
EAST_HEAD = "SimHei"
LATIN = "Times New Roman"
GRAY = RGBColor(0x59, 0x59, 0x59)
LIGHT = RGBColor(0x6E, 0x6E, 0x6E)
BLACK = RGBColor(0x14, 0x14, 0x14)
LANG = "zh"  # zh -> Chinese journal style; en -> English journal style


def _set_east(run, name):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def style_run(run, size=10.5, bold=False, italic=False, color=None, font=None):
    run.font.name = font or LATIN
    _set_east(run, EAST)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Italic handling: gene symbols and Latin binomials are italic, while
# protein/enzyme abbreviations stay regular.
# ---------------------------------------------------------------------------

PLANT_GENERA = {
    "Arabidopsis", "Antirrhinum", "Mimulus", "Glycine", "Fragaria",
    "Vaccinium", "Camellia", "Paeonia", "Dendrobium", "Phalaenopsis",
    "Oncidium", "Paphiopedilum", "Lilium", "Nicotiana", "Solanum", "Vitis",
    "Oryza", "Triticum", "Zea", "Gossypium", "Brassica", "Citrus", "Malus",
    "Rosa", "Acer", "Nelumbo", "Ipomoea", "Rhododendron", "Prunus", "Musa",
    "Actinidia", "Cymbidium", "Iris", "Osmanthus", "Anthurium", "Petunia",
    "Gentiana", "Delphinium", "Senecio", "Clitoria", "Phaseolus", "Capsicum",
    "Ginkgo", "Hordeum", "Pyrus", "Pisum", "Medicago", "Lotus", "Helianthus",
    "Aquilegia", "Nigella", "Coffea", "Pinus", "Abies", "Picea", "Magnolia",
    "Populus", "Eucalyptus", "Quercus", "Salix", "Ulmus", "Platanus",
    "Robinia", "Olea", "Elaeis", "Cocos", "Theobroma", "Rubus", "Sorghum",
    "Brachypodium", "Chrysanthemum", "Gerbera", "Tulipa", "Gladiolus",
    "Zantedeschia", "Caladium", "Zingiber", "Canna", "Heliconia",
    "Strelitzia", "Juglans", "Castanea", "Corylus", "Morus", "Ficus",
    "Cucumis", "Citrullus", "Cucurbita", "Daucus", "Allium", "Asparagus",
    "Physalis", "Atropa", "Ilex", "Hibiscus", "Raphanus", "Lactuca",
    "Artemisia", "Calendula", "Tagetes", "Zinnia", "Dahlia", "Aster",
    "Cosmos", "Centaurea", "Echinacea", "Rudbeckia", "Solidago", "Silybum",
    "Carthamus", "Linaria", "Saintpaulia", "Sinningia",
}

GENE_WORDS = {
    "NEGAN", "RTO", "YUP", "ROS1", "MYB", "MYB1", "MYB5", "MYB10", "MYB11",
    "MYB12", "MYB44", "MYBF", "MYBPA1.1", "MYBPA2.2", "MYC1", "MYC2",
    "bHLH1", "bHLH2", "bHLH3", "bHLH33", "MIANbHLH1", "WD40", "CHS", "CHI",
    "F3H", "F3'H", "F3'5'H", "DFR", "ANS", "UFGT", "LDOX", "LAR", "ANR",
    "FLS", "FNS", "TT2", "TT8", "TT12", "TTG1", "GL3", "EGL3", "AN1",
    "AN2", "AN11", "JAF13", "PAP1", "PAP2", "CYC", "CYC2a", "CYC2b",
    "DICH", "RAD", "DIV", "DRIF", "LEAFY", "CYCLOIDEA", "DICHOTOMA",
    "DIVARICATA", "RADIALIS", "APETALA1", "APETALA3", "PISTILLATA",
    "AGAMOUS", "SEPALLATA3", "WUSCHEL", "CLAVATA1", "SVP", "AGL6", "AGL24",
    "SOC1", "LFY", "WUS", "CLV1", "PTL", "CLV3", "HORT1", "GMYB10",
    "MdMYB1", "MdMYB10", "miR858",
}

_BINOMIAL = r"\b[A-Z][a-z]{2,}(?:\s*[×x]\s*|\s+)[a-z]{3,}\b"
_STOP_SPECIES = {
    "flower", "species", "plant", "plants", "gene", "genes", "hybrid",
    "variety", "group", "family", "class", "model", "review", "research",
    "analysis", "study", "studies", "development", "regulation", "protein",
    "proteins", "pathway", "complex", "biosynthesis", "fruit", "color",
    "colour", "flesh", "box", "domain", "floral", "homeotic", "meristem",
    "flowers", "organs", "organ", "seeds", "mutants", "varieties",
    "cultivars", "transcription", "factor", "factors", "expression",
    "receptor", "kinase", "synthase", "reductase", "transferase",
    "hydroxylase", "oxidase", "peroxidase", "activator", "repressor",
    "regulator", "regulators", "mutant", "callus", "culture", "cultures",
    "cells", "cell", "tissue", "tissues", "morphogenesis", "regeneration",
    "during", "using", "through", "with", "into", "from", "for", "the",
    "of", "in", "on", "at", "by", "to", "and", "response", "responses",
    "signal", "signaling", "signalling", "stress", "resistance",
    "tolerance", "activity", "activities", "function", "functions",
}
_PROT_CONN = r"[\s（(，,:：、与和及兼具活性]*"
_GENE_RES = [(w, re.compile(r"(?<![A-Za-z0-9])" + re.escape(w) + r"(?![A-Za-z0-9])"))
             for w in sorted(GENE_WORDS, key=len, reverse=True)]


def _protein_or_enzyme(text, s, e):
    before = text[max(0, s - 16):s]
    after = text[e:e + 14]
    if re.search(r"(?:蛋白|酶)" + _PROT_CONN + r"$", before):
        return True
    if re.search(r"^" + _PROT_CONN + r"(?:蛋白|酶)", after):
        return True
    return False


def italic_spans(text):
    spans = []
    for m in re.finditer(_BINOMIAL, text):
        seg = m.group(0)
        mm = re.match(r"^([A-Z][a-z]{2,})[\s×]+(.+)$", seg, re.S)
        if mm and mm.group(1) in PLANT_GENERA:
            if mm.group(2).lower() in _STOP_SPECIES:
                spans.append((m.start(), m.start() + len(mm.group(1))))
            else:
                spans.append((m.start(), m.end()))
    for word, pattern in _GENE_RES:
        for m in pattern.finditer(text):
            if not _protein_or_enzyme(text, m.start(), m.end()):
                spans.append((m.start(), m.end()))
    spans.sort()
    merged = []
    for s, e in spans:
        if merged and s <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], e))
        else:
            merged.append((s, e))
    return merged


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def _add_plain_runs(p, text, size, bold, color, italicize):
    if not text:
        return
    if not italicize:
        style_run(p.add_run(text), size=size, bold=bold, color=color)
        return
    pos = 0
    for s, e in italic_spans(text):
        if s > pos:
            style_run(p.add_run(text[pos:s]), size=size, bold=bold, color=color)
        style_run(p.add_run(text[s:e]), size=size, bold=bold, italic=True, color=color)
        pos = e
    if pos < len(text):
        style_run(p.add_run(text[pos:]), size=size, bold=bold, color=color)


def add_runs(p, text, size=10.5, bold=False, color=None, italicize=True):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            style_run(p.add_run(part[2:-2]), size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            style_run(p.add_run(part[1:-1]), size=size - 0.5, color=color, font="Consolas")
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            style_run(p.add_run(part[1:-1]), size=size, italic=True, color=color)
        else:
            _add_plain_runs(p, part, size, bold, color, italicize)


def add_paragraph(doc, text, size=10.5, bold=False, center=False, gray=False,
                  indent=False, after=3, italicize=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5 if LANG == "en" else 1.3
    p.paragraph_format.space_after = Pt(after)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Pt(36) if LANG == "en" else Pt(size * 2)
    color = LIGHT if gray else None
    add_runs(p, text, size=size, bold=bold, color=color, italicize=italicize)
    return p


def add_list_item(doc, text, numbered=False, num=None):
    if numbered:
        # Render the source number literally (1, 2, 3 ...) so each list keeps
        # its own sequential numbering instead of Word auto-continuing lists.
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5 if LANG == "en" else 1.3
        p.paragraph_format.space_after = Pt(3)
        li = 36 if LANG == "en" else 24
        p.paragraph_format.left_indent = Pt(li)
        p.paragraph_format.first_line_indent = Pt(-li)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_run(p.add_run(f"{num or 1}. "), size=12 if LANG == "en" else 10.5)
        add_runs(p, text)
        return p
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5 if LANG == "en" else 1.3
    p.paragraph_format.space_after = Pt(3)
    add_runs(p, text)
    return p


def add_reference(doc, text):
    """Hanging-indent reference entry (Chinese 8 pt / English journal 10 pt)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if LANG == "en":
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-36)
    else:
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.first_line_indent = Pt(-16)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text, size=10 if LANG == "en" else 8)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5 if LANG == "en" else 1.3
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        if LANG == "en":
            r.font.name = LATIN
            r.font.size = Pt(16)
            r.font.bold = True
        else:
            r.font.name = EAST_HEAD
            _set_east(r, EAST_HEAD)
            r.font.size = Pt(20)
        r.font.color.rgb = BLACK
        return p
    if level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        if LANG == "en":
            r.font.name = LATIN
            r.font.size = Pt(14)
            r.font.bold = True
        else:
            _set_east(r, EAST)
            r.font.size = Pt(14)
        r.font.color.rgb = BLACK
        return p
    p.paragraph_format.space_before = Pt(8 if level == 3 else 6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    if LANG == "en":
        r.font.name = LATIN
        r.font.size = Pt(12 if level == 3 else 12)
        r.font.bold = True
    else:
        r.font.name = EAST_HEAD
        _set_east(r, EAST_HEAD)
        r.font.size = Pt(11 if level == 3 else 10.5)
        r.font.bold = level >= 4
    r.font.color.rgb = BLACK
    return p


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            add_runs(p, val, size=10 if LANG == "en" else 9.5, bold=(i == 0), italicize=False)
    doc.add_paragraph()


def add_page_number(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def field(instr):
        r = p.add_run()
        r.font.size = Pt(9)
        _set_east(r, EAST)
        fld1 = OxmlElement("w:fldChar")
        fld1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        r._element.append(fld1)
        r._element.append(it)
        r._element.append(fld2)

    def text_run(s):
        r = p.add_run(s)
        r.font.size = Pt(9)
        _set_east(r, EAST)

    if LANG == "en":
        text_run("Page ")
        field(" PAGE ")
        text_run(" of ")
        field(" NUMPAGES ")
    else:
        text_run("第 ")
        field(" PAGE ")
        text_run(" 页 / 共 ")
        field(" NUMPAGES ")
        text_run(" 页")


def build(md_text, output):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(12 if LANG == "en" else 10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), EAST)
    normal.paragraph_format.line_spacing = 1.5 if LANG == "en" else 1.3

    lines = md_text.splitlines()
    i = 0
    in_refs = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not is_table_sep(lines[i]):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue

        if not stripped:
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).strip()
            add_heading(doc, title, level)
            if level <= 2:
                in_refs = "参考文献" in title
            i += 1
            continue

        if re.match(r"^---+$", stripped) or re.match(r"^\*\*\*+$", stripped):
            i += 1
            continue

        if stripped.startswith(">"):
            add_paragraph(doc, stripped.lstrip(">").strip(), gray=True,
                          size=(10 if LANG == "en" else 8.5) if in_refs else (10.5 if LANG == "en" else 9.5),
                          indent=False)
            i += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bullet:
            if in_refs:
                add_reference(doc, bullet.group(1).strip())
            else:
                add_list_item(doc, bullet.group(1).strip())
            i += 1
            continue

        numbered = re.match(r"^(\d+)[.、)]\s+(.*)$", stripped)
        if numbered:
            if in_refs:
                add_reference(doc, numbered.group(2).strip())
            else:
                add_list_item(doc, numbered.group(2).strip(), numbered=True,
                              num=int(numbered.group(1)))
            i += 1
            continue

        add_paragraph(doc, stripped,
                      size=(10 if LANG == "en" else 9) if in_refs else (12 if LANG == "en" else 10.5),
                      gray=in_refs,
                      indent=not in_refs)
        i += 1

    add_page_number(doc)
    doc.save(output)
    return output


def verify(output):
    check = Document(str(output))
    texts = [p.text for p in check.paragraphs]
    suspicious = set("\u9225\u604A\u94FF\u4E67\u788C\u4E6A\u4FF9\u50E3\u614F\u8133\u923C\u4E76\u617A\u4E7A\uE6A9\u20AC")
    bad_pat = re.compile("[\uE000-\uF8FF\uFB00-\uFB06]")
    with zipfile.ZipFile(str(output)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    hits = {ch for ch in xml if ch in suspicious or bad_pat.match(ch)}
    status = "通过" if not hits else "发现异常字符"
    print(f"校验：重新打开 {len(texts)} 个段落 | 乱码扫描 {status}")


def main():
    global LANG
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("inputs", nargs="+", help="Markdown 文件路径（多个时按顺序合并）")
    parser.add_argument("-o", "--output", default=None, help="输出 Word 路径（默认与输入同名 .docx）")
    parser.add_argument("--lang", choices=["zh", "en"], default="zh", help="排版语言：zh=中文期刊规范（默认），en=英文期刊格式")
    args = parser.parse_args()
    LANG = args.lang

    parts = []
    for name in args.inputs:
        input_path = Path(name)
        if not input_path.exists():
            sys.exit(f"输入文件不存在: {input_path}")
        parts.append(input_path.read_text(encoding="utf-8-sig"))
    md_text = "\n\n".join(parts)
    first = Path(args.inputs[0])
    output = Path(args.output) if args.output else first.with_suffix(".docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    build(md_text, output)
    print(f"已生成: {output}")
    verify(output)


if __name__ == "__main__":
    main()
